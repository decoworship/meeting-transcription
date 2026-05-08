"""Transcript export formats: SRT, VTT, DOCX, TXT."""

from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor

from ..transcription.base import TranscriptionResult


def _format_srt_time(seconds: float) -> str:
    """Format seconds as HH:MM:SS,mmm (SRT format)."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds - int(seconds)) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def _format_vtt_time(seconds: float) -> str:
    """Format seconds as HH:MM:SS.mmm (WebVTT format)."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds - int(seconds)) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"


def to_srt(result: TranscriptionResult, speaker_names: dict[str, str], include_speakers: bool) -> str:
    """Generate SRT subtitle format."""
    lines = []
    for i, seg in enumerate(result.segments, start=1):
        start = _format_srt_time(seg.start)
        end = _format_srt_time(seg.end)
        text = seg.text.strip()
        if include_speakers and seg.speaker:
            name = speaker_names.get(seg.speaker, seg.speaker)
            text = f"{name}: {text}"
        lines.append(f"{i}")
        lines.append(f"{start} --> {end}")
        lines.append(text)
        lines.append("")
    return "\n".join(lines)


def to_vtt(result: TranscriptionResult, speaker_names: dict[str, str], include_speakers: bool) -> str:
    """Generate WebVTT subtitle format."""
    lines = ["WEBVTT", ""]
    for i, seg in enumerate(result.segments, start=1):
        start = _format_vtt_time(seg.start)
        end = _format_vtt_time(seg.end)
        text = seg.text.strip()
        if include_speakers and seg.speaker:
            name = speaker_names.get(seg.speaker, seg.speaker)
            text = f"<v {name}>{text}"
        lines.append(f"{start} --> {end}")
        lines.append(text)
        lines.append("")
    return "\n".join(lines)


# Color palette for speakers in DOCX
SPEAKER_COLORS = [
    RGBColor(0x1F, 0x6A, 0xA5),  # blue
    RGBColor(0xD9, 0x53, 0x4F),  # red
    RGBColor(0x2E, 0xA8, 0x4F),  # green
    RGBColor(0xC4, 0x7F, 0x17),  # orange
    RGBColor(0x7B, 0x47, 0xA1),  # purple
    RGBColor(0x0E, 0x8B, 0x8B),  # teal
    RGBColor(0xB8, 0x4A, 0x7E),  # pink
]


def _format_time_label(seconds: float) -> str:
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def write_docx(
    result: TranscriptionResult,
    speaker_names: dict[str, str],
    include_speakers: bool,
    client: str,
    project: str,
    meeting_date: str,
    output_path: str,
) -> str:
    """Write transcription as a formatted Word document."""
    doc = Document()

    # Title
    title = doc.add_heading("Meeting Transcription", level=1)

    # Metadata
    meta_para = doc.add_paragraph()
    if client:
        meta_para.add_run(f"Client: ").bold = True
        meta_para.add_run(f"{client}\n")
    if project:
        meta_para.add_run(f"Project: ").bold = True
        meta_para.add_run(f"{project}\n")
    if meeting_date:
        meta_para.add_run(f"Date: ").bold = True
        meta_para.add_run(f"{meeting_date}\n")
    if result.duration:
        meta_para.add_run(f"Duration: ").bold = True
        meta_para.add_run(f"{_format_time_label(result.duration)}\n")
    if result.language:
        meta_para.add_run(f"Language: ").bold = True
        meta_para.add_run(f"{result.language}\n")

    if include_speakers:
        speakers = sorted(set(seg.speaker for seg in result.segments if seg.speaker and seg.speaker != "Unknown"))
        if speakers:
            meta_para.add_run(f"Speakers: ").bold = True
            meta_para.add_run(f"{len(speakers)}\n")
            for original in speakers:
                display = speaker_names.get(original, original)
                meta_para.add_run(f"  - {display}\n")

    doc.add_paragraph("_" * 60)

    # Build speaker -> color mapping
    speaker_color_map: dict[str, RGBColor] = {}
    if include_speakers:
        unique_speakers = sorted(set(seg.speaker for seg in result.segments if seg.speaker and seg.speaker != "Unknown"))
        for i, sp in enumerate(unique_speakers):
            speaker_color_map[sp] = SPEAKER_COLORS[i % len(SPEAKER_COLORS)]

    # Body
    for seg in result.segments:
        para = doc.add_paragraph()

        # Timestamp (gray, smaller)
        ts_run = para.add_run(f"[{_format_time_label(seg.start)}] ")
        ts_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        ts_run.font.size = Pt(9)

        # Speaker name (colored, bold)
        if include_speakers and seg.speaker:
            name = speaker_names.get(seg.speaker, seg.speaker)
            sp_run = para.add_run(f"{name}: ")
            sp_run.bold = True
            if seg.speaker in speaker_color_map:
                sp_run.font.color.rgb = speaker_color_map[seg.speaker]

        # Text
        para.add_run(seg.text.strip())

    doc.save(output_path)
    return output_path
